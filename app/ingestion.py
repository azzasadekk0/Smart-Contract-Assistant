from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import AppSettings


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def _extract_pdf_text(path: Path) -> str:
    try:
        import fitz

        with fitz.open(path) as doc:
            return "\n".join(page.get_text("text") for page in doc)
    except Exception:
        pass

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return "\n".join((page.extract_text() or "") for page in pdf.pages)
    except Exception:
        pass

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx_text(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def chunk_text(text: str, settings: AppSettings) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_text(text)


def ingest_files(
    paths: list[Path],
    settings: AppSettings,
    source_name_overrides: dict[str, str] | None = None,
) -> list[Document]:
    documents: list[Document] = []
    source_name_overrides = source_name_overrides or {}
    for path in paths:
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file extension: {suffix}")

        raw_text = extract_text(path).strip()
        if not raw_text:
            continue

        chunks = chunk_text(raw_text, settings)
        source_name = source_name_overrides.get(str(path), path.name)
        for idx, chunk in enumerate(chunks, start=1):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": source_name,
                        "chunk_id": idx,
                        "total_chunks": len(chunks),
                        "extension": suffix,
                    },
                )
            )
    return documents
